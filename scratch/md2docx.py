import docx
import sys

def md_to_docx(md_path, docx_path):
    doc = docx.Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    for line in lines:
        line = line.strip('\n')
        
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line)
            try:
                p.style = 'No Spacing'
                if p.runs:
                    p.runs[0].font.name = 'Courier New'
            except:
                pass
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('---'):
            pass # ignore hr
        elif line.strip() == '':
            pass
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            doc.add_paragraph(line[3:], style='List Number')
        else:
            doc.add_paragraph(line)
            
    doc.save(docx_path)

if __name__ == '__main__':
    md_to_docx(sys.argv[1], sys.argv[2])
